from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import pandas as pd
from datetime import datetime

class MeetingDocumentProcessor:
    def __init__(self):
        self.doc = None
        self.content_updates = []
        
    def read_document(self, file_path):
        """读取会议文档"""
        try:
            self.doc = Document(file_path)
            return True
        except Exception as e:
            print(f"读取文档失败: {str(e)}")
            return False
            
    def add_content_update(self, section, original_text, updated_text_cn, updated_text_en):
        """添加内容更新记录"""
        self.content_updates.append({
            'section': section,
            'original_text': original_text,
            'updated_text_cn': updated_text_cn,
            'updated_text_en': updated_text_en
        })
        
    def generate_update_report(self, output_path):
        """生成更新报告"""
        if not self.content_updates:
            return False
            
        # 创建新的Word文档
        report_doc = Document()
        report_doc.add_heading('文档更新报告', 0)
        report_doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # 添加更新内容表格
        table = report_doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '章节'
        header_cells[1].text = '原文'
        header_cells[2].text = '更新内容（中文）'
        header_cells[3].text = '更新内容（英文）'
        
        # 添加更新内容
        for update in self.content_updates:
            row_cells = table.add_row().cells
            row_cells[0].text = update['section']
            row_cells[1].text = update['original_text']
            row_cells[2].text = update['updated_text_cn']
            row_cells[3].text = update['updated_text_en']
            
        # 保存报告
        try:
            report_doc.save(output_path)
            return True
        except Exception as e:
            print(f"保存报告失败: {str(e)}")
            return False
            
    def update_document(self, output_path):
        """更新文档内容，保持原有格式"""
        if not self.doc:
            return False
            
        # 创建新文档
        new_doc = Document()
        
        # 复制原文档内容并保持格式
        for para in self.doc.paragraphs:
            new_para = new_doc.add_paragraph()
            # 复制段落格式
            new_para.style = para.style
            new_para.paragraph_format.alignment = para.paragraph_format.alignment
            new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
            new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            
            # 复制文本和格式
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                new_run.font.color.rgb = run.font.color.rgb
                
        # 复制表格
        for table in self.doc.tables:
            new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = table.style
            
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.cell(i, j)
                    new_cell.text = cell.text
                    # 复制单元格格式
                    new_cell.paragraphs[0].paragraph_format.alignment = cell.paragraphs[0].paragraph_format.alignment
                    
        # 保存更新后的文档
        try:
            new_doc.save(output_path)
            return True
        except Exception as e:
            print(f"保存文档失败: {str(e)}")
            return False
            
    def export_to_excel(self, output_path):
        """导出更新内容到Excel"""
        if not self.content_updates:
            return False
            
        df = pd.DataFrame(self.content_updates)
        try:
            df.to_excel(output_path, index=False)
            return True
        except Exception as e:
            print(f"导出Excel失败: {str(e)}")
            return False 