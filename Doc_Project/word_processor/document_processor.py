from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

class DocumentProcessor:
    def __init__(self):
        self.doc = None

    def read_document(self, file_path):
        """读取Word文档"""
        try:
            self.doc = Document(file_path)
            return True
        except Exception as e:
            print(f"读取文档失败: {str(e)}")
            return False

    def get_document_content(self):
        """获取文档内容"""
        if not self.doc:
            return "未加载文档"
        
        content = []
        for para in self.doc.paragraphs:
            content.append(para.text)
        return '\n'.join(content)

    def optimize_content(self):
        """优化文档内容"""
        if not self.doc:
            return False

        # 移除多余的空行
        for para in self.doc.paragraphs:
            if not para.text.strip():
                p = para._element
                p.getparent().remove(p)

        # 统一段落格式
        for para in self.doc.paragraphs:
            if para.text.strip():
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in para.runs:
                    run.font.size = Pt(12)
                    run.font.name = '宋体'

        return True

    def add_summary(self, summary_text):
        """添加文档摘要"""
        if not self.doc:
            return False

        # 在文档开头添加摘要
        summary_para = self.doc.add_paragraph()
        summary_para.add_run("文档摘要：").bold = True
        summary_para.add_run(f"\n{summary_text}")
        return True

    def save_document(self, file_path):
        """保存文档"""
        if not self.doc:
            return False
        try:
            self.doc.save(file_path)
            return True
        except Exception as e:
            print(f"保存文档失败: {str(e)}")
            return False

    def extract_keywords(self):
        """提取文档关键词"""
        if not self.doc:
            return []

        # 简单的关键词提取（这里可以集成更复杂的NLP算法）
        text = self.get_document_content()
        words = re.findall(r'\w+', text)
        word_freq = {}
        for word in words:
            if len(word) > 1:  # 忽略单字词
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # 返回出现频率最高的10个词
        return sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]

    def add_watermark(self, watermark_text):
        """添加水印"""
        if not self.doc:
            return False

        # 在文档末尾添加水印
        watermark = self.doc.add_paragraph()
        watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = watermark.add_run(watermark_text)
        run.font.color.rgb = RGBColor(200, 200, 200)  # 浅灰色
        return True 