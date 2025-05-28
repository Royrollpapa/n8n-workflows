"""
会议处理器模块
提供会议文档的处理功能
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class MeetingProcessor:
    """会议文档处理器类"""
    
    def __init__(self):
        """初始化会议处理器"""
        pass
        
    def extract_agenda(self, file_path):
        """提取会议议程
        
        Args:
            file_path: 会议文档路径
            
        Returns:
            议程文本
        """
        doc = Document(file_path)
        agenda = []
        
        # 查找议程部分
        found_agenda = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "议程" in text or "会议议程" in text:
                found_agenda = True
                continue
            if found_agenda and text:
                if "决议" in text or "会议纪要" in text:
                    break
                agenda.append(text)
                
        return "\n".join(agenda)
        
    def extract_resolutions(self, file_path):
        """提取会议决议
        
        Args:
            file_path: 会议文档路径
            
        Returns:
            决议文本
        """
        doc = Document(file_path)
        resolutions = []
        
        # 查找决议部分
        found_resolutions = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "决议" in text or "会议决议" in text:
                found_resolutions = True
                continue
            if found_resolutions and text:
                if "会议纪要" in text:
                    break
                resolutions.append(text)
                
        return "\n".join(resolutions)
        
    def generate_summary(self, file_path, meeting_info):
        """生成会议纪要
        
        Args:
            file_path: 会议文档路径
            meeting_info: 会议信息字典，包含title、date、attendees等
            
        Returns:
            会议纪要Document对象
        """
        doc = Document()
        
        # 添加标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(meeting_info["title"])
        title_run.font.size = Pt(16)
        title_run.bold = True
        
        # 添加会议信息
        info = doc.add_paragraph()
        info.add_run(f"会议时间：{meeting_info['date']}\n")
        info.add_run(f"参会人员：{', '.join(meeting_info['attendees'])}\n")
        
        # 添加议程
        agenda = self.extract_agenda(file_path)
        if agenda:
            doc.add_paragraph("会议议程：").bold = True
            doc.add_paragraph(agenda)
            
        # 添加决议
        resolutions = self.extract_resolutions(file_path)
        if resolutions:
            doc.add_paragraph("会议决议：").bold = True
            doc.add_paragraph(resolutions)
            
        return doc
        
    def save_summary(self, doc, file_path):
        """保存会议纪要
        
        Args:
            doc: Document对象
            file_path: 保存路径
        """
        doc.save(file_path) 