"""
Markdown文档处理器
处理Markdown文档的转换和格式化
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any
import markdown
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt
from .presentation_processor import PresentationProcessor

class MarkdownProcessor:
    """Markdown文档处理器类"""
    
    def __init__(self):
        self.ppt_processor = PresentationProcessor()
        
    def to_html(self, input_path: str, output_path: str, 
                template_path: Optional[str] = None,
                **kwargs) -> str:
        """
        将Markdown转换为HTML
        
        Args:
            input_path: 输入Markdown文件路径
            output_path: 输出HTML文件路径
            template_path: HTML模板路径（可选）
            **kwargs: 其他参数
            
        Returns:
            str: 输出文件路径
        """
        # 读取Markdown文件
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
            
        # 转换为HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # 如果提供了模板，使用模板包装内容
        if template_path and os.path.exists(template_path):
            with open(template_path, 'r', encoding='utf-8') as f:
                template = f.read()
            html_content = template.replace('{{content}}', html_content)
            
        # 保存HTML文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
            
        return output_path
    
    def to_ppt(self, input_path: str, output_path: str,
               template_path: Optional[str] = None,
               **kwargs) -> str:
        """
        将Markdown转换为PowerPoint
        
        Args:
            input_path: 输入Markdown文件路径
            output_path: 输出PPT文件路径
            template_path: PPT模板路径（可选）
            **kwargs: 其他参数
            
        Returns:
            str: 输出文件路径
        """
        # 先转换为HTML
        temp_html = str(Path(output_path).with_suffix('.html'))
        self.to_html(input_path, temp_html)
        
        # 从HTML转换为PPT
        result = self.ppt_processor.from_html(temp_html, output_path, template_path)
        
        # 清理临时文件
        if os.path.exists(temp_html):
            os.remove(temp_html)
            
        return result
    
    def to_docx(self, input_path: str, output_path: str,
                template_path: Optional[str] = None,
                **kwargs) -> str:
        """
        将Markdown转换为Word文档
        
        Args:
            input_path: 输入Markdown文件路径
            output_path: 输出Word文件路径
            template_path: Word模板路径（可选）
            **kwargs: 其他参数
            
        Returns:
            str: 输出文件路径
        """
        # 创建新的Word文档
        doc = docx.Document(template_path) if template_path else docx.Document()
        
        # 读取Markdown文件
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
            
        # 转换为HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # 解析HTML并添加到Word文档
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                paragraph = doc.add_paragraph(element.get_text())
                paragraph.style = f'Heading {level}'
            elif element.name == 'p':
                doc.add_paragraph(element.get_text())
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
            elif element.name == 'table':
                table = doc.add_table(rows=1, cols=len(element.find_all('th')))
                table.style = 'Table Grid'
                
                # 添加表头
                header_cells = table.rows[0].cells
                for i, th in enumerate(element.find_all('th')):
                    header_cells[i].text = th.get_text()
                
                # 添加数据行
                for tr in element.find_all('tr')[1:]:
                    row_cells = table.add_row().cells
                    for i, td in enumerate(tr.find_all('td')):
                        row_cells[i].text = td.get_text()
        
        # 保存文档
        doc.save(output_path)
        return output_path 