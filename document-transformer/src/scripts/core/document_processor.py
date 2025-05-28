"""
文档处理器模块
提供Word文档的基本处理功能
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import nltk
from nltk.tokenize import word_tokenize
from collections import Counter

class DocumentProcessor:
    """Word文档处理器类"""
    
    def __init__(self):
        """初始化文档处理器"""
        # 下载NLTK数据
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt')
            
    def read_document(self, file_path):
        """读取Word文档
        
        Args:
            file_path: Word文档路径
            
        Returns:
            Document对象
        """
        return Document(file_path)
        
    def save_document(self, doc, file_path):
        """保存Word文档
        
        Args:
            doc: Document对象
            file_path: 保存路径
        """
        doc.save(file_path)
        
    def optimize_format(self, doc):
        """优化文档格式
        
        Args:
            doc: Document对象
            
        Returns:
            优化后的Document对象
        """
        # 统一字体和大小
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'
                run.font.size = Pt(12)
                
        # 统一段落对齐方式
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        return doc
        
    def add_summary(self, doc, summary):
        """添加文档摘要
        
        Args:
            doc: Document对象
            summary: 摘要文本
            
        Returns:
            添加摘要后的Document对象
        """
        # 在文档开头添加摘要
        summary_paragraph = doc.add_paragraph()
        summary_paragraph.add_run("摘要：").bold = True
        summary_paragraph.add_run(summary)
        
        return doc
        
    def extract_keywords(self, doc, top_n=10):
        """提取文档关键词
        
        Args:
            doc: Document对象
            top_n: 返回的关键词数量
            
        Returns:
            关键词列表，每个元素为(词, 频次)元组
        """
        # 提取所有文本
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + " "
            
        # 分词并统计频次
        words = word_tokenize(text.lower())
        word_freq = Counter(words)
        
        # 返回频次最高的词
        return word_freq.most_common(top_n)
        
    def add_watermark(self, doc, watermark_text):
        """添加文档水印
        
        Args:
            doc: Document对象
            watermark_text: 水印文本
            
        Returns:
            添加水印后的Document对象
        """
        # 在每页添加水印
        for section in doc.sections:
            header = section.header
            paragraph = header.paragraphs[0]
            run = paragraph.add_run(watermark_text)
            run.font.color.rgb = RGBColor(200, 200, 200)  # 浅灰色
            run.font.size = Pt(36)
            
        return doc 